"""
document_preprocessor.py
--------------------------
Pre-processing pipeline that takes raw documents from the Roche Drive,
adds metadata, converts them to clean markdown, and writes the result
to a "processed" output location.

WHY THIS EXISTS:
  Nachi identified that retrieval quality is limited by having to work
  against raw DOCX/PDF formatting. This module solves that by producing
  a clean, markdown version of every document with metadata headers,
  stored separately from the original Roche source documents.

SCOPE — IMPORTANT:
  This module NEVER modifies or deletes the original Roche documents.
  Read access to originals stays read-only at all times.

KNOWN PLATFORM CONSTRAINT — WRITE TARGET (read before touching this file):
  Google Drive service accounts have no storage quota of their own and
  cannot write into a personal-Drive-owned folder, even with Editor
  permissions — Google requires either a genuine Shared Drive or OAuth
  user delegation for any service-account write. Our IE student Google
  Workspace does not permit creating Shared Drives, so writing the
  processed output back to Drive is not available in this environment.

  DEV/DEMO FALLBACK (current implementation):
    Processed markdown is written to a LOCAL folder (data/processed/)
    and committed to the repo. This is safe ONLY because all current
    source documents are synthetic/mock — never real Roche data.

  PRODUCTION PATH (not implemented here, documented for the report):
    A real deployment would require either:
      (a) a genuine Google Shared Drive as the write target, or
      (b) OAuth user delegation (a human-authenticated token, not a
          service account) so writes are attributed to a real user's
          quota.
    Processed output derived from REAL Roche documents must never be
    committed to source control or written outside Roche-controlled
    infrastructure. The local-write fallback below is explicitly a
    placeholder for non-confidential synthetic data only and must be
    replaced before any real-data deployment.

OUTPUT FORMAT — example processed file:

    ---
    source_file: "New Employee Onboarding & Digital Access Guide.docx"
    source_id: "1-nIDDnfQ2MAWSneCEoYc6VE5endh315-"
    department: "IT / Onboarding"
    process_type: "onboarding"
    processed_at: "2026-06-29T10:00:00Z"
    ---

    # New Employee Onboarding & Digital Access Guide

    ## Welcome to the Basel Research Team!

    As a new researcher or technician at the Basel campus...

    ## 1. Mandatory Core Applications & Permissions Table

    | Application Name | Purpose | Approval Required By | Where to Request |
    |---|---|---|---|
    | ChemVantage | Ordering lab chemicals... | Lab Manager | UAP -> CHEM-VAN-01 |

HOW TO RUN:
    python document_preprocessor.py

  Reads from GOOGLE_SERVICE_ACCOUNT_JSON + DRIVE_FOLDER_ID in .env
  (same credentials your GoogleDriveSource already uses).
  Writes one .md file per source document into data/processed/ locally.
"""

from __future__ import annotations

import io
import logging
import os
import re
from dataclasses import dataclass
from datetime import datetime, timezone
from typing import Optional

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Scopes — kept separate and explicit, matching Nachi's existing pattern.
# Read is used for listing/downloading originals. Write is used ONLY
# for creating the "processed" subfolder and uploading files into it.
# ---------------------------------------------------------------------------
_READ_SCOPES  = ["https://www.googleapis.com/auth/drive.readonly"]
_WRITE_SCOPES = ["https://www.googleapis.com/auth/drive"]

_SUPPORTED_MIME = {
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
    "application/pdf": "pdf",
    "application/vnd.google-apps.document": "gdoc",
    "text/plain": "text",
    "text/markdown": "text",
}

PROCESSED_FOLDER_NAME = "processed"

# Simple keyword-based department/process tagging.
# Extend this dict as more document types are added — it's intentionally
# simple for the MVP rather than another ML classification step.
_DEPARTMENT_KEYWORDS = {
    "onboarding":      ["onboarding", "new employee", "access guide", "digital access"],
    "waste":           ["waste", "disposal", "decontamination"],
    "procurement":     ["ordering", "chemicals", "consumables", "purchasing"],
    "logistics":       ["return", "conveyor", "shipping", "material return"],
    "lab_operations":  ["calibration", "instrument", "cleaning", "equipment"],
    "facilities":      ["campus", "building", "access, dining", "facilities"],
    "collaboration":   ["sharing", "workflows", "collaborative"],
}


@dataclass
class ProcessedDocument:
    source_id:     str
    source_name:   str
    department:    str
    process_type:  str
    markdown:      str

    def to_file_content(self) -> str:
        """Builds the final .md file content with YAML frontmatter metadata."""
        timestamp = datetime.now(timezone.utc).isoformat()
        frontmatter = (
            "---\n"
            f'source_file: "{self.source_name}"\n'
            f'source_id: "{self.source_id}"\n'
            f'department: "{self.department}"\n'
            f'process_type: "{self.process_type}"\n'
            f'processed_at: "{timestamp}"\n'
            "---\n\n"
        )
        return frontmatter + self.markdown


# ---------------------------------------------------------------------------
# Tagging logic
# ---------------------------------------------------------------------------

def infer_department_and_process(filename: str, content: str) -> tuple[str, str]:
    """
    Lightweight keyword match against filename + content to assign a
    department/process tag. Falls back to 'general' if nothing matches.
    Intentionally simple — a rules-based first pass, easy to extend later.
    """
    haystack = (filename + " " + content[:500]).lower()

    for process_type, keywords in _DEPARTMENT_KEYWORDS.items():
        if any(kw in haystack for kw in keywords):
            department = process_type.replace("_", " ").title()
            return department, process_type

    return "General", "general"


# ---------------------------------------------------------------------------
# Content -> Markdown conversion
# ---------------------------------------------------------------------------

def docx_to_markdown(content: bytes) -> str:
    """
    Converts a DOCX file's bytes into clean markdown.
    Headings become '#'/'##', tables become markdown tables,
    paragraphs stay as plain text.
    """
    import docx
    document = docx.Document(io.BytesIO(content))
    lines: list[str] = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading 1" in style or "title" in style:
            lines.append(f"# {text}")
        elif "heading 2" in style:
            lines.append(f"## {text}")
        elif "heading 3" in style:
            lines.append(f"### {text}")
        else:
            lines.append(text)
        lines.append("")  # blank line between blocks

    for table in document.tables:
        if not table.rows:
            continue
        header = [cell.text.strip() for cell in table.rows[0].cells]
        lines.append("| " + " | ".join(header) + " |")
        lines.append("|" + "|".join(["---"] * len(header)) + "|")
        for row in table.rows[1:]:
            cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
        lines.append("")

    return "\n".join(lines)


def pdf_to_markdown(content: bytes) -> str:
    """Converts PDF text extraction into loosely-formatted markdown.
    PDFs don't carry structural info the way DOCX does, so this is a
    straightforward paragraph dump with light cleanup."""
    import pypdf
    reader = pypdf.PdfReader(io.BytesIO(content))
    raw_text = "\n".join(page.extract_text() or "" for page in reader.pages)

    # Collapse excessive blank lines, keep paragraph breaks
    cleaned = re.sub(r"\n{3,}", "\n\n", raw_text)
    return cleaned.strip()


def plain_text_to_markdown(content: bytes) -> str:
    return content.decode("utf-8", errors="replace")


# ---------------------------------------------------------------------------
# Main preprocessing pipeline
# ---------------------------------------------------------------------------

class DocumentPreprocessor:
    """
    Orchestrates: list raw docs -> convert to markdown -> tag metadata
    -> write processed output.

    Read operations use _READ_SCOPES against the real Google Drive.

    WRITE TARGET: see the module docstring's "KNOWN PLATFORM CONSTRAINT"
    section. Service-account writes to a personal-Drive folder are not
    possible in this environment, so output is written locally via
    `_write_local_fallback()` — explicitly named to flag that this is a
    dev/demo substitute, not the production write path. Real Drive
    write support (_get_write_service, kept below) is retained for
    when a Shared Drive or OAuth-delegated credential becomes
    available, but is not currently invoked by `run()`.
    """

    def __init__(
        self,
        folder_id: str,
        local_output_dir: str = "data/processed",
        credentials_path: Optional[str] = None,
    ):
        self.folder_id = folder_id
        self.local_output_dir = local_output_dir
        self.credentials_path = credentials_path or os.environ["GOOGLE_SERVICE_ACCOUNT_JSON"]
        self._read_service  = None
        self._write_service = None
        self._processed_folder_id: Optional[str] = None

    # ---- service builders, explicitly separate per scope ----

    def _get_read_service(self):
        if self._read_service:
            return self._read_service
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
        creds = service_account.Credentials.from_service_account_file(
            self.credentials_path, scopes=_READ_SCOPES
        )
        self._read_service = build("drive", "v3", credentials=creds, cache_discovery=False)
        return self._read_service

    def _get_write_service(self):
        """Kept for the production Drive-write path (Shared Drive / OAuth
        delegation). Not called by run() today — see module docstring."""
        if self._write_service:
            return self._write_service
        from google.oauth2 import service_account
        from googleapiclient.discovery import build
        creds = service_account.Credentials.from_service_account_file(
            self.credentials_path, scopes=_WRITE_SCOPES
        )
        self._write_service = build("drive", "v3", credentials=creds, cache_discovery=False)
        return self._write_service

    # ---- step 1: ensure processed/ subfolder exists ----

    def _ensure_processed_folder(self) -> str:
        """Finds or creates the 'processed' subfolder. Returns its folder ID."""
        if self._processed_folder_id:
            return self._processed_folder_id

        read_service = self._get_read_service()
        resp = read_service.files().list(
            q=(
                f"'{self.folder_id}' in parents and trashed=false "
                f"and name='{PROCESSED_FOLDER_NAME}' "
                f"and mimeType='application/vnd.google-apps.folder'"
            ),
            fields="files(id, name)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()

        existing = resp.get("files", [])
        if existing:
            self._processed_folder_id = existing[0]["id"]
            logger.info("preprocessor.folder.found", extra={"folder_id": self._processed_folder_id})
            return self._processed_folder_id

        # Doesn't exist yet — create it using the WRITE service.
        # supportsAllDrives=True is required so the new folder is created
        # INSIDE the Shared Drive (using its pooled storage) rather than
        # attributed to the service account's own (zero) storage quota.
        write_service = self._get_write_service()
        try:
            folder = write_service.files().create(
                body={
                    "name": PROCESSED_FOLDER_NAME,
                    "mimeType": "application/vnd.google-apps.folder",
                    "parents": [self.folder_id],
                },
                supportsAllDrives=True,
                fields="id",
            ).execute()
        except Exception as e:
            if "storageQuotaExceeded" in str(e):
                raise RuntimeError(
                    "storageQuotaExceeded: the service account has no personal Drive "
                    "storage. This folder must live inside a Shared Drive (not 'My Drive') "
                    "for supportsAllDrives=True to use the Shared Drive's pooled storage. "
                    "Confirm DRIVE_FOLDER_ID points to a folder inside a Shared Drive, "
                    "not a personal Drive folder shared with the service account."
                ) from e
            raise

        self._processed_folder_id = folder["id"]
        logger.info("preprocessor.folder.created", extra={"folder_id": self._processed_folder_id})
        return self._processed_folder_id

    # ---- step 2: list + convert each source document ----

    def _list_source_documents(self) -> list[dict]:
        """Lists documents in the root folder, excluding the processed/ subfolder itself."""
        read_service = self._get_read_service()
        resp = read_service.files().list(
            q=f"'{self.folder_id}' in parents and trashed=false",
            fields="files(id, name, mimeType)",
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()

        files = [
            f for f in resp.get("files", [])
            if f["mimeType"] in _SUPPORTED_MIME
        ]
        return files

    def _download(self, file_id: str, mime: str) -> bytes:
        read_service = self._get_read_service()
        if mime == "application/vnd.google-apps.document":
            return read_service.files().export_media(
                fileId=file_id, mimeType="text/plain"
            ).execute()

        from googleapiclient.http import MediaIoBaseDownload
        request = read_service.files().get_media(fileId=file_id, supportsAllDrives=True)
        buf = io.BytesIO()
        downloader = MediaIoBaseDownload(buf, request)
        done = False
        while not done:
            _, done = downloader.next_chunk()
        return buf.getvalue()

    def _convert_to_markdown(self, raw: bytes, mime: str) -> str:
        if mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            return docx_to_markdown(raw)
        elif mime == "application/pdf":
            return pdf_to_markdown(raw)
        else:
            return plain_text_to_markdown(raw)

    def process_one(self, file_meta: dict) -> ProcessedDocument:
        raw = self._download(file_meta["id"], file_meta["mimeType"])
        markdown_body = self._convert_to_markdown(raw, file_meta["mimeType"])
        department, process_type = infer_department_and_process(file_meta["name"], markdown_body)

        return ProcessedDocument(
            source_id=file_meta["id"],
            source_name=file_meta["name"],
            department=department,
            process_type=process_type,
            markdown=markdown_body,
        )

    # ---- step 3: write processed output ----
    #
    # DEV/DEMO FALLBACK: writes locally instead of back to Drive.
    # See the module docstring's "KNOWN PLATFORM CONSTRAINT" section for why.
    # Safe only because current source documents are synthetic/mock data.
    # A production deployment with real Roche documents must replace this
    # with a write to a genuine Shared Drive or an OAuth-delegated upload —
    # never local disk or source control for real data.

    def _write_local_fallback(self, doc: ProcessedDocument) -> str:
        os.makedirs(self.local_output_dir, exist_ok=True)

        safe_name = re.sub(r"[^\w\-. ]", "_", doc.source_name)
        safe_name = re.sub(r"\.(docx|pdf|txt)$", "", safe_name, flags=re.IGNORECASE)
        output_name = f"{safe_name}.md"
        output_path = os.path.join(self.local_output_dir, output_name)

        with open(output_path, "w", encoding="utf-8") as f:
            f.write(doc.to_file_content())

        logger.info("preprocessor.file.written_local", extra={"file_path": output_path})
        return output_path

    # ---- public entry point ----

    def run(self) -> list[dict]:
        """
        Runs the full pipeline: list -> convert -> tag -> write.
        Returns a summary list of what was processed.
        Output currently goes to local disk — see _write_local_fallback.
        """
        sources = self._list_source_documents()
        logger.info("preprocessor.run.start", extra={"count": len(sources)})

        results = []
        for file_meta in sources:
            try:
                doc = self.process_one(file_meta)
                output_path = self._write_local_fallback(doc)
                results.append({
                    "source_name":  doc.source_name,
                    "department":   doc.department,
                    "process_type": doc.process_type,
                    "output_path":  output_path,
                    "status":       "success",
                })
            except Exception as e:
                logger.error("preprocessor.file.failed", extra={"file_name": file_meta["name"], "error": str(e)})
                results.append({
                    "source_name": file_meta["name"],
                    "status":      "failed",
                    "error":       str(e),
                })

        return results


# ---------------------------------------------------------------------------
# CLI entry point
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    import json
    from dotenv import load_dotenv

    logging.basicConfig(level=logging.INFO, format="%(asctime)s %(message)s")
    load_dotenv()

    folder_id = os.environ["DRIVE_FOLDER_ID"]
    preprocessor = DocumentPreprocessor(
        folder_id=folder_id,
        local_output_dir="data/processed",
    )

    print(f"Starting preprocessing for folder {folder_id}...")
    results = preprocessor.run()

    print("\n" + "=" * 60)
    print("PREPROCESSING SUMMARY")
    print("=" * 60)
    for r in results:
        status_icon = "✅" if r["status"] == "success" else "❌"
        print(f"{status_icon} {r['source_name']}")
        if r["status"] == "success":
            print(f"   department: {r['department']} | process_type: {r['process_type']}")
        else:
            print(f"   error: {r['error']}")

    success_count = sum(1 for r in results if r["status"] == "success")
    print(f"\n{success_count}/{len(results)} documents processed successfully.")
    print(f"Check the '{preprocessor.local_output_dir}/' folder for output.")
    print(
        "\nNote: this writes locally (dev/demo fallback for synthetic data only). "
        "See module docstring for the production write-path requirement."
    )
