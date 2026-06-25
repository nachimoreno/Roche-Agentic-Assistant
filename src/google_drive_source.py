"""
google_drive_source.py
----------------------
A `DocumentSource` implementation that reads from Google Drive.
Drops in as a direct replacement for `LocalMarkdownSource` — same
`SourceDocument` output shape, same `list_documents()` interface.

HOW TO PLUG IN:
    Set `DOCUMENT_SOURCE=google_drive` in .env — `main.build_source()` then
    constructs this from `Settings`. To wire it manually somewhere else:
        from google_drive_source import GoogleDriveSource
        source = GoogleDriveSource(folder_id=settings.drive_folder_id)
    # pass `source` wherever LocalMarkdownSource was passed — nothing else changes.

SETUP:
    1. Create a Google Cloud project + enable the Drive API.
    2. Create a Service Account, download the JSON key.
    3. Share your Drive folder with the service account email (Viewer role).
    4. Add to .env:
         GOOGLE_SERVICE_ACCOUNT_JSON=path/to/key.json
         DRIVE_FOLDER_ID=<your folder id from the Drive URL>

AUTH FALLBACK (for local dev without a service account):
    Set GOOGLE_OAUTH_CREDENTIALS=path/to/oauth_token.json instead — a token
    file produced by the standard google-auth-oauthlib InstalledAppFlow.

SUPPORTED FILE TYPES:
    - Google Docs        → exported as plain text
    - Google Sheets      → exported as CSV text
    - .txt / .md files   → downloaded directly
    - PDFs               → text extracted (requires pypdf)
    - .docx              → text extracted (requires python-docx)

VERSION DEDUPLICATION:
    If multiple files in the folder share the same name, only the most
    recently modified one is returned. This handles Roche's known problem
    of duplicate/outdated document versions living side by side in Drive.

SHAREPOINT MIGRATION PATH:
    When Roche moves to SharePoint, implement a `SharePointSource` class
    with the same `list_documents()` -> Iterator[SourceDocument] signature.
    Nothing in retrieval.py, agent.py, or orchestrator.py needs to change.
"""

from __future__ import annotations

import io
import json
import logging
import os
from dataclasses import dataclass, field
from datetime import datetime, timezone
from typing import Any, Iterator

logger = logging.getLogger(__name__)

# ---------------------------------------------------------------------------
# Re-use the SourceDocument dataclass already defined in document_source.py.
# Import it from there so the rest of the pipeline sees one type, not two.
# ---------------------------------------------------------------------------
try:
    from document_source import SourceDocument
except ImportError:
    # Standalone fallback if running this file directly for testing
    @dataclass
    class SourceDocument:
        id: str
        title: str
        content: str
        modified_at: datetime
        metadata: dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Supported Google mime types and how to handle them
# ---------------------------------------------------------------------------
_GDOC_EXPORT_MIME   = "text/plain"
_GSHEET_EXPORT_MIME = "text/csv"

# Read covers listing/downloading the corpus; write is needed only to upload a
# new document into the folder (the "add document" feature). The write scope is
# requested lazily, on a separate service, so normal ingestion stays read-only.
_READ_SCOPES  = ["https://www.googleapis.com/auth/drive.readonly"]
_WRITE_SCOPES = ["https://www.googleapis.com/auth/drive"]

# Extensions a scientist may upload, mapped to the MIME type Drive should store
# them as. Google-native types (Docs/Sheets) are intentionally absent — those
# are created in Drive, not uploaded.
_UPLOAD_MIME = {
    ".pdf":  "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt":  "text/plain",
    ".md":   "text/markdown",
}

_SUPPORTED = {
    "application/vnd.google-apps.document":    "gdoc",
    "application/vnd.google-apps.spreadsheet": "gsheet",
    "text/plain":                               "text",
    "text/markdown":                            "text",
    "application/pdf":                          "pdf",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
}


class GoogleDriveSource:
    """
    DocumentSource that lists and downloads files from a Google Drive folder.

    Parameters
    ----------
    folder_id:
        The Drive folder ID (from the URL: drive.google.com/drive/folders/<ID>).
    credentials_path:
        Path to a service account JSON key or OAuth2 token file.
        Falls back to GOOGLE_SERVICE_ACCOUNT_JSON / GOOGLE_OAUTH_CREDENTIALS env vars.
    recursive:
        If True, walk sub-folders too. Default True — Roche docs are spread
        across subfolders so this should almost always stay True.
    """

    def __init__(
        self,
        folder_id: str,
        credentials_path: str | None = None,
        recursive: bool = True,
    ) -> None:
        self.folder_id = folder_id
        self.recursive = recursive
        self._creds_path = credentials_path or os.getenv(
            "GOOGLE_SERVICE_ACCOUNT_JSON",
            os.getenv("GOOGLE_OAUTH_CREDENTIALS"),
        )
        self._service = None         # lazy read-only service
        self._writer_service = None  # lazy read/write service (uploads only)

    # ------------------------------------------------------------------
    # DocumentSource protocol
    # ------------------------------------------------------------------

    def list_documents(self) -> Iterator[SourceDocument]:
        """
        Yield one SourceDocument per file in the Drive folder.
        Deduplicates by name — newest modified_at wins.
        """
        raw = self._list_files(self.folder_id)
        deduped = _deduplicate(raw)

        logger.info(
            "drive.documents.listed",
            extra={"folder": self.folder_id, "count": len(deduped)},
        )

        for meta in deduped:
            content = self._download_text(meta)
            if not content.strip():
                logger.warning("drive.document.empty", extra={"doc_name": meta["name"]})
                continue

            modified_at = _parse_time(meta.get("modifiedTime", ""))
            title = _infer_title(meta["name"], content)

            yield SourceDocument(
                id=meta["id"],
                title=title,
                content=content,
                modified_at=modified_at,
                metadata={
                    "drive_id":      meta["id"],
                    "drive_name":    meta["name"],
                    "mime_type":     meta.get("mimeType", ""),
                    "modified_time": meta.get("modifiedTime", ""),
                    "source":        "google_drive",
                    "folder_id":     self.folder_id,
                    "subfolder":     meta.get("subfolder", ""),
                    # Deep link back to the file in Drive, so the UI can turn a
                    # citation into a click-through to the source document.
                    "url":           meta.get("webViewLink", ""),
                },
            )

    def check_connection(self) -> int:
        """Cheaply verify the integration is live, for startup health reporting.

        Loads credentials and lists a single page of the target folder. Returns
        the number of files visible on that page (a lower bound, not a full
        count). Raises on any failure — missing/invalid credentials, an
        inaccessible folder, or a network error — so callers can surface *why*.
        """
        service = self._get_service()
        resp = service.files().list(
            q=f"'{self.folder_id}' in parents and trashed=false",
            fields="files(id)",
            pageSize=10,
            supportsAllDrives=True,
            includeItemsFromAllDrives=True,
        ).execute()
        return len(resp.get("files", []))

    # ------------------------------------------------------------------
    # Internal helpers
    # ------------------------------------------------------------------

    def _get_service(self):
        """Lazy-init the read-only Drive API v3 service."""
        if self._service:
            return self._service

        from googleapiclient.discovery import build

        creds = self._load_credentials(_READ_SCOPES)
        self._service = build("drive", "v3", credentials=creds, cache_discovery=False)
        return self._service

    def _get_writer_service(self):
        """Lazy-init a read/write Drive service, used only for uploads."""
        if self._writer_service:
            return self._writer_service

        from googleapiclient.discovery import build

        creds = self._load_credentials(_WRITE_SCOPES)
        self._writer_service = build(
            "drive", "v3", credentials=creds, cache_discovery=False
        )
        return self._writer_service

    def upload_file(
        self, *, filename: str, data: bytes, mime_type: str | None = None
    ) -> "SourceDocument":
        """Upload a file into the configured Drive folder and return it as a
        `SourceDocument` (id, title, extracted text, metadata) ready to ingest.

        The text is extracted from the bytes we already hold, so there's no
        download round-trip. Raises on any Drive error (bad scope, no Editor
        access, or the service-account storage-quota limit on My Drive folders)
        — the caller surfaces the reason.
        """
        from googleapiclient.http import MediaIoBaseUpload

        ext = os.path.splitext(filename)[1].lower()
        mime = mime_type or _UPLOAD_MIME.get(ext)
        if mime is None:
            raise ValueError(f"Unsupported file type: {ext or filename!r}")

        service = self._get_writer_service()
        media = MediaIoBaseUpload(io.BytesIO(data), mimetype=mime, resumable=False)
        meta = service.files().create(
            body={"name": filename, "parents": [self.folder_id]},
            media_body=media,
            fields="id, name, mimeType, modifiedTime, webViewLink",
            supportsAllDrives=True,
        ).execute()

        content = extract_text(filename, data, mime)
        if not content.strip():
            raise ValueError(
                "Could not extract any text from the document "
                "(it may be a scanned image or empty)."
            )

        return SourceDocument(
            id=meta["id"],
            title=_infer_title(meta["name"], content),
            content=content,
            modified_at=_parse_time(meta.get("modifiedTime", "")),
            metadata={
                "drive_id":      meta["id"],
                "drive_name":    meta["name"],
                "mime_type":     meta.get("mimeType", mime),
                "modified_time": meta.get("modifiedTime", ""),
                "source":        "google_drive",
                "folder_id":     self.folder_id,
                "subfolder":     "",
                "url":           meta.get("webViewLink", ""),
            },
        )

    def can_write(self) -> bool:
        """Best-effort probe: can this account actually upload into the folder?

        Does a *non-destructive* metadata read of the target folder's
        `capabilities.canAddChildren` through the write-scoped service, so the
        answer reflects what a real upload would be allowed to do (it catches a
        viewer-only service account, which is the common misconfiguration). Never
        raises — returns False on any failure (missing creds, no access, network)
        so callers can grey out the upload feature without risking startup.
        """
        try:
            service = self._get_writer_service()
            meta = service.files().get(
                fileId=self.folder_id,
                fields="capabilities/canAddChildren",
                supportsAllDrives=True,
            ).execute()
            return bool(meta.get("capabilities", {}).get("canAddChildren"))
        except Exception:
            logger.warning("drive.can_write_probe_failed", exc_info=True)
            return False

    def _load_credentials(self, scopes: list[str]):
        if not self._creds_path:
            raise RuntimeError(
                "No Google credentials found. "
                "Set GOOGLE_SERVICE_ACCOUNT_JSON or GOOGLE_OAUTH_CREDENTIALS."
            )

        # The credential is accepted two ways:
        #   * a filesystem path to a key file (local dev), or
        #   * the raw JSON content itself (hosted deploys like HF Spaces, where
        #     the key is injected as a secret env var, not a committed file).
        # A leading "{" means inline JSON; anything else is treated as a path.
        value = self._creds_path.strip()
        raw = value if value.startswith("{") else open(value).read()
        info = json.loads(raw)

        if info.get("type") == "service_account":
            from google.oauth2 import service_account
            return service_account.Credentials.from_service_account_info(
                info, scopes=scopes
            )
        else:
            from google.oauth2.credentials import Credentials
            return Credentials.from_authorized_user_info(info, scopes)

    def _list_files(self, folder_id: str, subfolder_name: str = "") -> list[dict]:
        """Return flat list of file metadata dicts from a Drive folder.
        supportsAllDrives + includeItemsFromAllDrives are required for
        Shared Drives — without them the API returns nothing."""
        service = self._get_service()
        files: list[dict] = []
        page_token = None

        while True:
            resp = service.files().list(
                q=f"'{folder_id}' in parents and trashed=false",
                fields="nextPageToken, files(id, name, mimeType, modifiedTime, webViewLink)",
                pageToken=page_token,
                pageSize=100,
                supportsAllDrives=True,
                includeItemsFromAllDrives=True,
            ).execute()

            for f in resp.get("files", []):
                mime = f.get("mimeType", "")
                if mime == "application/vnd.google-apps.folder":
                    if self.recursive:
                        files.extend(self._list_files(f["id"], subfolder_name=f["name"]))
                elif mime in _SUPPORTED:
                    f["subfolder"] = subfolder_name  # track which subfolder this came from
                    files.append(f)
                else:
                    logger.debug(
                        "drive.file.skipped",
                        extra={"doc_name": f["name"], "mime": mime},
                    )

            page_token = resp.get("nextPageToken")
            if not page_token:
                break

        return files

    def _download_text(self, meta: dict) -> str:
        """Download a Drive file and return its text content."""
        service = self._get_service()
        mime = meta.get("mimeType", "")
        file_id = meta["id"]

        try:
            if mime == "application/vnd.google-apps.document":
                return _export(service, file_id, _GDOC_EXPORT_MIME)

            elif mime == "application/vnd.google-apps.spreadsheet":
                return _export(service, file_id, _GSHEET_EXPORT_MIME)

            elif mime == "application/pdf":
                raw = _download_bytes(service, file_id)
                return _pdf_to_text(raw)

            elif mime == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                raw = _download_bytes(service, file_id)
                return _docx_to_text(raw)

            else:  # text/plain, text/markdown
                raw = _download_bytes(service, file_id)
                return raw.decode("utf-8", errors="replace")

        except Exception as exc:
            logger.error(
                "drive.download.failed",
                extra={"doc_name": meta["name"], "error": str(exc)},
            )
            return ""


# ---------------------------------------------------------------------------
# Module-level helpers (no state, easy to test)
# ---------------------------------------------------------------------------

def _export(service, file_id: str, mime_type: str) -> str:
    data = service.files().export_media(
        fileId=file_id,
        mimeType=mime_type,
        supportsAllDrives=True,
    ).execute()
    return data.decode("utf-8", errors="replace")


def _download_bytes(service, file_id: str) -> bytes:
    from googleapiclient.http import MediaIoBaseDownload
    request = service.files().get_media(fileId=file_id, supportsAllDrives=True)
    buf = io.BytesIO()
    dl = MediaIoBaseDownload(buf, request)
    done = False
    while not done:
        _, done = dl.next_chunk()
    return buf.getvalue()


def _pdf_to_text(content: bytes) -> str:
    try:
        import pypdf
        reader = pypdf.PdfReader(io.BytesIO(content))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except ImportError:
        logger.warning("pypdf not installed — PDF text extraction unavailable")
        return "[PDF: install pypdf to extract text]"


def _docx_to_text(content: bytes) -> str:
    try:
        import docx

        document = docx.Document(io.BytesIO(content))
        parts = []

        # Regular paragraphs
        for p in document.paragraphs:
            text = p.text.strip()
            if text:
                parts.append(text)

        # Tables
        for table in document.tables:
            for row in table.rows:
                cells = [
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                ]
                if cells:
                    parts.append(" | ".join(cells))

        return "\n".join(parts)

    except ImportError:
        logger.warning("python-docx not installed — DOCX text extraction unavailable")
        return "[DOCX: install python-docx to extract text]"

def extract_text(filename: str, data: bytes, mime_type: str = "") -> str:
    """Extract plain text from uploaded bytes, dispatching on extension/MIME.

    Mirrors the per-type handling in `_download_text`, but works on bytes we
    already have in memory (an upload) instead of a Drive download.
    """
    ext = os.path.splitext(filename)[1].lower()
    if ext == ".pdf" or mime_type == "application/pdf":
        return _pdf_to_text(data)
    if ext == ".docx" or mime_type.endswith("wordprocessingml.document"):
        return _docx_to_text(data)
    # text/plain, text/markdown, .txt, .md
    return data.decode("utf-8", errors="replace")


def _deduplicate(files: list[dict]) -> list[dict]:
    """
    Keep only the most recently modified file when multiple share a name.
    Handles Roche's known Drive version-drift problem.
    """
    best: dict[str, dict] = {}
    for f in files:
        key = f["name"].lower().strip()
        if key not in best:
            best[key] = f
        else:
            if f.get("modifiedTime", "") > best[key].get("modifiedTime", ""):
                logger.info(
                    "drive.dedup",
                    extra={
                        "kept_name": f["name"],
                        "dropped_time": best[key].get("modifiedTime"),
                        "kept_time": f.get("modifiedTime"),
                    },
                )
                best[key] = f
    return list(best.values())


def _parse_time(iso: str) -> datetime:
    if not iso:
        return datetime.now(tz=timezone.utc)
    try:
        return datetime.fromisoformat(iso.replace("Z", "+00:00"))
    except ValueError:
        return datetime.now(tz=timezone.utc)


def _infer_title(filename: str, content: str) -> str:
    """Use the first Markdown H1 if present, else the filename without extension."""
    for line in content.splitlines():
        if line.startswith("# "):
            return line[2:].strip()
    # Strip known extensions
    for ext in (".md", ".txt", ".pdf", ".csv"):
        if filename.lower().endswith(ext):
            return filename[: -len(ext)]
    return filename
